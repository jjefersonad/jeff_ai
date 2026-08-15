"""Testes do DocxWriter (infrastructure/docx_writer) e da tool create_docx_document.

Cobrem os critérios de aceitação da task `custom-office-doc-tools-task-docx-1`:
- REQ-001: gera `.docx` válido só com python-docx, sem pandoc/soffice/node.
- REQ-002: suporta headings, parágrafos, listas (ordenada/não), tabela e imagem.
- REQ-003: salva em `outputs/documents/docx/<ts>.docx`, retorna
  `{path, url, metadata}` e a URL começa com `/api/files/docx/`.
- REQ-005: entrada inválida retorna erro descritivo e não deixa arquivo parcial.
- Tool como adapter fino + wiring em `composition.dependencies`.

E os da task `fix-docx-generation-url-and-json-content-task-test-1`:
- REQ-005: `create_docx_document` retorna `url` absoluta (default e com
  `BASE_URL` sobrescrita).
- REQ-006: payload string serializado em JSON produz blocos (não título com
  JSON bruto); JSON malformado retorna `{"error": ...}` sem arquivo parcial.

E os da task `fix-docx-empty-content-task-docx-2`:
- REQ-006 (MODIFIED): string simples não-JSON deixou de ser um "modo legado"
  que produzia documento só com título — agora é rejeitada com `DomainError`
  (causa raiz de um bug onde documentos gerados só tinham título). A
  tolerância a JSON serializado como string continua funcionando sem mudança.
"""
from __future__ import annotations

import json
import re
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document as DocxReader

import src.composition.dependencies as dep
import src.tools.create_docx_document_tool as docx_tool
from src.domain.documents import (
    DocxSpec,
    Heading,
    ImageRef,
    ListBlock,
    Paragraph,
    Table,
)
from src.domain.shared.errors import DomainError
from src.infrastructure.documents.docx_writer import DocxWriter
from src.models.docx_document import DocxBlockInput, DocxDocumentInput

# --- helpers ---------------------------------------------------------------


def _make_png(path: Path, width: int = 2, height: int = 2) -> str:
    """Grava um PNG mínimo válido (2x2) no caminho e retorna o caminho como str.

    Evita dependência de Pillow nos testes — gera bytes PNG manualmente
    (magic + IHDR + IDAT com pixels zero + IEND).
    """
    def _chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 0, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\x00\x00\x00\x00" * width for _ in range(height))
    idat = zlib.compress(raw)
    png_bytes = sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")
    path.write_bytes(png_bytes)
    return str(path)


# --- DocxWriter (infra) ----------------------------------------------------


@pytest.fixture
def writer(tmp_path: Path) -> DocxWriter:
    """DocxWriter apontando para um diretório temporário (sem poluir outputs/)."""
    return DocxWriter(output_dir=tmp_path, url_prefix="/api/files")


async def test_writer_creates_valid_docx_with_headings_and_table(writer, tmp_path):
    """REQ-001/REQ-002/REQ-003: writer gera .docx válido e estrutura conforme spec."""
    spec = DocxSpec(
        title="Relatório",
        blocks=(
            Heading(text="Seção 1", level=1),
            Paragraph(text="Parágrafo introdutório."),
            ListBlock(items=("Item A", "Item B"), ordered=False),
            Table(
                rows=(("Col 1", "Col 2"), ("v1", "v2")),
                header=True,
            ),
        ),
    )

    result = await writer.write(spec)

    # REQ-003: contrato {path, url, metadata}, URL pública e arquivo no disco.
    assert result.path
    assert result.url.startswith("/api/files/docx/")
    assert result.url.endswith(".docx")
    assert result.metadata == {"kind": "docx", "title": "Relatório", "block_count": 4}
    path = Path(result.path)
    assert path.is_file()
    assert path.parent == tmp_path

    # Reabre o arquivo com python-docx e valida a estrutura.
    doc = DocxReader(str(path))
    # Título do documento: estilo "Title". Headings: "Heading 1".."Heading 9".
    titles = [p.text for p in doc.paragraphs if p.style.name == "Title"]
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading ")]
    assert titles == ["Relatório"]
    assert headings == ["Seção 1"]

    # Tabela: 2 colunas x 2 linhas, cabeçalho em negrito.
    tables = doc.tables
    assert len(tables) == 1
    table = tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert table.cell(0, 0).text == "Col 1"
    assert table.cell(1, 1).text == "v2"
    header_run = table.cell(0, 0).paragraphs[0].runs[0]
    assert header_run.bold is True


async def test_writer_sets_core_properties_title(writer):
    """fix-docx-core-properties-title: sem isso, clientes que leem metadata OOXML
    (ex.: preview de documento no WhatsApp) mostram 'Sem título'/'Untitled' —
    reportado ao vivo em `fix-whatsapp-document-delivery` mesmo com o `fileName`
    correto no payload de `send_media`, porque o título exibido vem do metadata
    interno do arquivo, não do nome do arquivo."""
    spec = DocxSpec(title="Ferramentas Disponíveis", blocks=(Paragraph(text="Lista."),))

    result = await writer.write(spec)

    doc = DocxReader(result.path)
    assert doc.core_properties.title == "Ferramentas Disponíveis"


async def test_writer_supports_ordered_list_and_image(writer, tmp_path):
    """REQ-002: lista numerada e inserção de imagem a partir de disco."""
    image_path = tmp_path / "logo.png"
    _make_png(image_path, width=4, height=4)

    spec = DocxSpec(
        title="Com imagem",
        blocks=(
            ListBlock(items=("Primeiro", "Segundo", "Terceiro"), ordered=True),
            ImageRef(path=str(image_path), width_inches=2.0),
        ),
    )

    result = await writer.write(spec)
    doc = DocxReader(result.path)

    # Lista numerada: estilo "List Number".
    numbered = [
        p.text
        for p in doc.paragraphs
        if p.style.name == "List Number"
    ]
    assert numbered == ["Primeiro", "Segundo", "Terceiro"]

    # Imagem embutida como parte do documento.
    inline_shapes = doc.inline_shapes
    assert len(inline_shapes) == 1


async def test_writer_missing_image_raises_without_partial_file(writer, tmp_path):
    """REQ-005: imagem inexistente levanta erro claro e nada fica no diretório."""
    spec = DocxSpec(
        title="Com imagem faltando",
        blocks=(ImageRef(path="/caminho/que/nao/existe.png"),),
    )

    with pytest.raises(RuntimeError, match="Imagem não encontrada"):
        await writer.write(spec)

    # Sem arquivo parcial: o diretório pode existir, mas deve estar vazio.
    assert list(tmp_path.iterdir()) == []


async def test_writer_rejects_non_docx_spec():
    """O writer só aceita DocxSpec — typecheck em tempo de execução."""
    writer = DocxWriter()
    with pytest.raises(TypeError, match="DocxSpec"):
        await writer.write("not a spec")  # type: ignore[arg-type]


# --- Domínio: validações que protegem o writer ----------------------------


def test_docx_spec_rejects_empty_title():
    with pytest.raises(DomainError, match="title"):
        DocxSpec(title="")


def test_heading_rejects_invalid_level():
    with pytest.raises(DomainError, match="level"):
        Heading(text="x", level=0)


def test_table_rejects_non_rectangular_rows():
    with pytest.raises(DomainError, match="retangular"):
        Table(rows=(("a", "b"), ("c",)))


def test_list_block_rejects_empty_items():
    with pytest.raises(DomainError, match="itens vazios"):
        ListBlock(items=("", "ok"))


def test_docx_spec_rejects_empty_blocks():
    """REQ-007: `blocks` vazio é barrado no domínio, antes do writer rodar."""
    with pytest.raises(DomainError, match="blocks"):
        DocxSpec(title="Relatório", blocks=())


def test_docx_spec_accepts_non_empty_blocks():
    """REQ-007 (regressão): `blocks` com ao menos um item continua funcionando."""
    spec = DocxSpec(title="Relatório", blocks=(Heading(text="Resumo", level=1),))
    assert spec.blocks == (Heading(text="Resumo", level=1),)


# --- Tool create_docx_document (HTML pipeline) -----------------------------


def _point_docx_tool(monkeypatch, tmp_path: Path) -> None:
    from unittest.mock import AsyncMock

    monkeypatch.setattr(docx_tool, "require_user_docs_dir", AsyncMock(return_value=tmp_path))
    monkeypatch.setattr(docx_tool, "_document_url_prefix", lambda: "/api/files")
    monkeypatch.setattr(docx_tool, "record_ownership", AsyncMock())


def test_to_blocks_from_structured_input():
    blocks = docx_tool._to_blocks(
        [
            DocxBlockInput(type="heading", text="H1", level=2),
            DocxBlockInput(type="paragraph", text="p"),
            DocxBlockInput(type="list", items=["a", "b"], ordered=True),
            DocxBlockInput(
                type="table",
                rows=[["c1", "c2"], ["v1", "v2"]],
                header=False,
            ),
            DocxBlockInput(type="image", path="/tmp/x.png", width_inches=1.5),
            DocxBlockInput(type="unknown", text="ignored"),
        ]
    )
    assert len(blocks) == 5  # type "unknown" é ignorado (contrato tolerante)
    assert isinstance(blocks[0], Heading)
    assert blocks[0].level == 2
    assert isinstance(blocks[2], ListBlock)
    assert blocks[2].ordered is True
    assert isinstance(blocks[3], Table)
    assert blocks[3].header is False
    assert isinstance(blocks[4], ImageRef)
    assert blocks[4].width_inches == 1.5


def test_to_blocks_drops_incomplete_blocks():
    """Bloco sem campos obrigatórios é descartado (não derruba a geração)."""
    blocks = docx_tool._to_blocks(
        [
            DocxBlockInput(type="heading", text=""),  # sem text → drop
            DocxBlockInput(type="list", items=[]),  # sem items → drop
            DocxBlockInput(type="image", path=""),  # sem path → drop
            DocxBlockInput(type="paragraph", text="ok"),
        ]
    )
    assert len(blocks) == 1
    assert isinstance(blocks[0], Paragraph)


async def test_tool_returns_path_url_metadata(monkeypatch, tmp_path):
    """A tool usa o pipeline HTML e devolve o contrato esperado."""
    _point_docx_tool(monkeypatch, tmp_path)

    result = await docx_tool.create_docx_document.coroutine(
        DocxDocumentInput(
            title="Hello",
            blocks=[DocxBlockInput(type="heading", text="Intro", level=1)],
        )
    )

    assert set(result) == {"path", "url", "metadata"}
    assert "/api/files/docx/" in result["url"]
    assert result["url"].endswith(".docx")
    assert result["metadata"]["kind"] == "docx"
    assert result["metadata"]["title"] == "Hello"
    assert Path(result["path"]).is_file()


async def test_tool_returns_error_on_invalid_input(monkeypatch, tmp_path):
    """REQ-005: entrada inválida vira `error` descritivo, sem arquivo parcial."""
    _point_docx_tool(monkeypatch, tmp_path)

    # Título vazio é barrado pelo domínio antes de qualquer I/O.
    result = await docx_tool.create_docx_document.coroutine(
        DocxDocumentInput(
            title="   ",
            blocks=[DocxBlockInput(type="paragraph", text="corpo")],
        )
    )

    assert "error" in result
    assert "title" in result["error"].lower()
    assert list(tmp_path.iterdir()) == []


# --- REQ-005 / root-env-config REQ-003: URL absoluta via BASE_URL -----------


@pytest.fixture
def _redirect_default_output_dir(monkeypatch, tmp_path):
    import src.infrastructure.documents.output_target as output_target

    monkeypatch.setattr(output_target, "_DEFAULT_BASE", tmp_path)


async def test_build_create_document_defaults_to_localhost_3000(
    monkeypatch, _redirect_default_output_dir
):
    monkeypatch.delenv("BASE_URL", raising=False)
    monkeypatch.delenv("FRONTEND_ORIGIN", raising=False)
    use_case = dep.build_create_document()

    result = await use_case.execute(
        DocxSpec(title="Só título", blocks=(Paragraph(text="corpo"),))
    )

    assert result.url.startswith("http://localhost:3000/api/files/docx/")


async def test_build_create_document_uses_base_url(
    monkeypatch, _redirect_default_output_dir
):
    """unify-root-env-and-compose-layout-task-code-1-unit-1: origem = BASE_URL."""
    monkeypatch.setenv("BASE_URL", "https://api.example.com")
    monkeypatch.delenv("FRONTEND_ORIGIN", raising=False)
    use_case = dep.build_create_document()

    result = await use_case.execute(
        DocxSpec(title="Só título", blocks=(Paragraph(text="corpo"),))
    )

    assert result.url.startswith("https://api.example.com/api/files/docx/")


def test_suite_does_not_monkeypatch_document_base_url():
    """unify-root-env-and-compose-layout-task-code-2-unit-1: suite uses BASE_URL only."""
    # Nome legado partido para este próprio arquivo não se auto-flagrar.
    # Só flagra uso ativo (setenv/delenv/getenv/environ) — asserts negativas
    # em testes de Compose (`"…" not in env`) podem citar o nome aposentado.
    legacy = "DOCUMENT" + "_BASE_URL"
    active = re.compile(
        rf'(?:monkeypatch\.(?:setenv|delenv)|os\.getenv|os\.environ(?:\.get)?)\(\s*["\']{re.escape(legacy)}["\']'
    )
    tests_root = Path(__file__).resolve().parent
    offenders = [
        str(path.relative_to(tests_root))
        for path in tests_root.rglob("*.py")
        if path.resolve() != Path(__file__).resolve()
        and active.search(path.read_text(encoding="utf-8"))
    ]
    assert offenders == [], f"{legacy} ainda usado ativamente em: {offenders}"


async def test_tool_returns_absolute_url_end_to_end(monkeypatch, tmp_path):
    """A tool `create_docx_document`, de ponta a ponta, devolve `url` absoluta."""
    from unittest.mock import AsyncMock

    monkeypatch.delenv("BASE_URL", raising=False)
    monkeypatch.delenv("FRONTEND_ORIGIN", raising=False)
    monkeypatch.setattr(docx_tool, "require_user_docs_dir", AsyncMock(return_value=tmp_path))
    monkeypatch.setattr(docx_tool, "record_ownership", AsyncMock())

    result = await docx_tool.create_docx_document.coroutine(
        DocxDocumentInput(
            title="Só título",
            blocks=[DocxBlockInput(type="paragraph", text="corpo")],
        )
    )

    assert result["url"].startswith("http://localhost:3000/api/files/docx/")


# --- REQ-006: string JSON serializada vs. string simples vs. JSON malformado


def test_parse_payload_from_json_serialized_string():
    """String que é, na verdade, um DocxDocumentInput serializado em JSON."""
    from src.application.documents.resolve_html_document_input import parse_tool_payload

    payload = json.dumps(
        {
            "title": "Relatório",
            "blocks": [
                {"type": "heading", "text": "Seção 1", "level": 1},
                {"type": "paragraph", "text": "Corpo do texto."},
            ],
        }
    )

    parsed = parse_tool_payload(payload)

    assert parsed.title == "Relatório"
    assert len(parsed.blocks or []) == 2
    assert parsed.blocks[0].type == "heading"
    assert parsed.blocks[1].type == "paragraph"


def test_parse_payload_rejects_plain_non_json_string():
    """String comum (não JSON) é rejeitada — não produz mais documento só com título."""
    from src.application.documents.resolve_html_document_input import parse_tool_payload

    with pytest.raises(DomainError, match="estruturado"):
        parse_tool_payload("Relatório trimestral")


async def test_tool_plain_non_json_string_returns_error_without_partial_file(
    monkeypatch, tmp_path
):
    """REQ-006 (MODIFIED): string simples não-JSON vira `{"error": ...}` via a tool."""
    _point_docx_tool(monkeypatch, tmp_path)

    result = await docx_tool.create_docx_document.coroutine("Relatório trimestral")

    assert "error" in result
    assert list(tmp_path.iterdir()) == []


def test_parse_payload_rejects_malformed_json_string():
    """`{`-prefixado mas JSON inválido levanta DomainError (não vira título)."""
    from src.application.documents.resolve_html_document_input import parse_tool_payload

    with pytest.raises(DomainError, match="JSON"):
        parse_tool_payload("{not valid json")


def test_parse_payload_rejects_json_shaped_but_schema_invalid():
    """`{`-prefixado, JSON válido, mas sem modo de entrada válido."""
    from src.application.documents.resolve_html_document_input import (
        parse_tool_payload,
        resolve_html_document_input,
    )

    parsed = parse_tool_payload(json.dumps({"blocks": []}))
    with pytest.raises(DomainError, match="vazia|blocks|obrigatório"):
        resolve_html_document_input(parsed)


async def test_tool_json_string_payload_produces_blocks_not_raw_json_title(
    monkeypatch, tmp_path
):
    """Regressão do bug: JSON serializado não pode virar título literal."""
    _point_docx_tool(monkeypatch, tmp_path)
    payload = json.dumps(
        {"title": "Doc via JSON", "blocks": [{"type": "paragraph", "text": "corpo"}]}
    )

    result = await docx_tool.create_docx_document.coroutine(payload)

    assert result["metadata"]["title"] == "Doc via JSON"
    assert result["metadata"]["kind"] == "docx"
    doc = DocxReader(result["path"])
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "Doc via JSON" in body
    assert "corpo" in body
    # O JSON bruto não aparece em nenhum lugar do documento.
    assert not any("{" in p.text for p in doc.paragraphs)


async def test_tool_malformed_json_string_returns_error_without_partial_file(
    monkeypatch, tmp_path
):
    _point_docx_tool(monkeypatch, tmp_path)

    result = await docx_tool.create_docx_document.coroutine("{not valid json")

    assert "error" in result
    assert list(tmp_path.iterdir()) == []
