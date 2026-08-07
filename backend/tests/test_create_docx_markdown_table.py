"""Tool-level coverage for Markdown-table rejection (fix-docx-table-markdown)."""
from __future__ import annotations

from pathlib import Path

import src.composition.dependencies as dep
import src.tools.create_docx_document_tool as docx_tool
from docx import Document as DocxReader
from src.infrastructure.documents.docx_writer import DocxWriter
from src.models.docx_document import DocxBlockInput, DocxDocumentInput

_MD_TABLE = "| A | B |\n|---|---|\n| 1 | 2 |"


async def test_tool_rejects_markdown_table_paragraph(monkeypatch, tmp_path):
    monkeypatch.setattr(
        docx_tool,
        "build_create_document",
        lambda: dep.CreateDocument(writer=DocxWriter(output_dir=tmp_path)),
    )

    result = await docx_tool.create_docx_document.coroutine(
        DocxDocumentInput(
            title="Relatório",
            blocks=[DocxBlockInput(type="paragraph", text=_MD_TABLE)],
        )
    )

    assert "error" in result
    assert list(tmp_path.iterdir()) == []


async def test_tool_accepts_prose_and_native_table(monkeypatch, tmp_path):
    monkeypatch.setattr(
        docx_tool,
        "build_create_document",
        lambda: dep.CreateDocument(writer=DocxWriter(output_dir=tmp_path)),
    )

    result = await docx_tool.create_docx_document.coroutine(
        DocxDocumentInput(
            title="Relatório",
            blocks=[
                DocxBlockInput(type="paragraph", text="Resumo do trimestre."),
                DocxBlockInput(
                    type="table",
                    rows=[["Mês", "Receita"], ["Jan", "12000"]],
                    header=True,
                ),
            ],
        )
    )

    assert "error" not in result
    assert "url" in result
    assert result["url"]
    assert Path(result["path"]).is_file()


async def test_tool_generated_docx_has_native_table_cells(monkeypatch, tmp_path):
    monkeypatch.setattr(
        docx_tool,
        "build_create_document",
        lambda: dep.CreateDocument(writer=DocxWriter(output_dir=tmp_path)),
    )

    rows = [["Mês", "Receita"], ["Jan", "12000"]]
    result = await docx_tool.create_docx_document.coroutine(
        DocxDocumentInput(
            title="Tabela",
            blocks=[
                DocxBlockInput(type="paragraph", text="Dados:"),
                DocxBlockInput(type="table", rows=rows, header=True),
            ],
        )
    )

    doc = DocxReader(result["path"])
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Mês"
    assert table.cell(0, 1).text == "Receita"
    assert table.cell(1, 0).text == "Jan"
    assert table.cell(1, 1).text == "12000"
    # Não deve haver pipes de markdown no corpo como parágrafo de tabela.
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "| Mês | Receita |" not in body_text


async def test_tool_error_mentions_type_table_and_rows(monkeypatch, tmp_path):
    monkeypatch.setattr(
        docx_tool,
        "build_create_document",
        lambda: dep.CreateDocument(writer=DocxWriter(output_dir=tmp_path)),
    )

    result = await docx_tool.create_docx_document.coroutine(
        DocxDocumentInput(
            title="Relatório",
            blocks=[DocxBlockInput(type="paragraph", text=_MD_TABLE)],
        )
    )

    err = result["error"]
    assert 'type="table"' in err or "type='table'" in err
    assert "rows" in err
