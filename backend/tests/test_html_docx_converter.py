"""Testes do converter HTML→DOCX (html-document-tools-task-docx-spike-1)."""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from src.infrastructure.documents.html_docx_converter import HtmlDocxConverter


_FIXTURE_HTML = """
<!DOCTYPE html>
<html>
<body>
  <h1>Título Principal</h1>
  <p>Introdução.</p>
  <table>
    <tr><th>A</th><th>B</th></tr>
    <tr><td>1</td><td>2</td></tr>
  </table>
</body>
</html>
"""


@pytest.mark.asyncio
async def test_html_heading_and_table_become_native_docx() -> None:
    """Unit-1: HTML com h1 + table 2x2 → heading e tabela nativa no .docx."""
    converter = HtmlDocxConverter()
    payload = await converter.convert(html=_FIXTURE_HTML, css=None, kind="docx")

    assert payload[:2] == b"PK"  # zip/docx
    doc = Document(BytesIO(payload))

    heading_texts = [
        p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")
    ]
    assert "Título Principal" in heading_texts

    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert table.cell(0, 0).text.strip() == "A"
    assert table.cell(0, 1).text.strip() == "B"
    assert table.cell(1, 0).text.strip() == "1"
    assert table.cell(1, 1).text.strip() == "2"


@pytest.mark.asyncio
async def test_full_document_with_head_meta_still_converts() -> None:
    """Regression: void tags em <head> não devem deixar skip_depth preso."""
    html = (
        '<!DOCTYPE html><html><head><meta charset="utf-8">'
        "<title>Relatório</title></head><body>"
        "<h1>Relatório</h1><p>Resumo do trimestre.</p>"
        "<table><tr><th>Mês</th><th>Receita</th></tr>"
        "<tr><td>Jan</td><td>12000</td></tr></table>"
        "</body></html>"
    )
    converter = HtmlDocxConverter()
    payload = await converter.convert(html=html, css=None, kind="docx")
    doc = Document(BytesIO(payload))
    texts = [p.text for p in doc.paragraphs]
    assert any("Relatório" in t for t in texts)
    assert any("Resumo" in t for t in texts)
    assert len(doc.tables) == 1


@pytest.mark.asyncio
async def test_converter_rejects_non_docx_kind() -> None:
    converter = HtmlDocxConverter()
    with pytest.raises(Exception, match="docx"):
        await converter.convert(html="<p>x</p>", css=None, kind="pdf")
