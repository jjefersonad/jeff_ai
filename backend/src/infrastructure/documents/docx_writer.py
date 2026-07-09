"""Adapter de criação de documentos Word via python-docx (implementa o port)."""
from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document
from docx.shared import Inches

from src.application.ports.document_writer import DocumentWriterPort
from src.domain.documents import (
    DocumentResult,
    DocumentSpec,
    DocxSpec,
    Heading,
    ImageRef,
    ListBlock,
    Paragraph,
    Table,
)
from src.infrastructure.documents.output_target import DocumentOutput


class DocxWriter(DocumentWriterPort):
    """Gera um `.docx` a partir de um `DocxSpec` usando apenas python-docx.

    Não invoca pandoc/soffice/node. Persiste em `outputs/documents/docx/` e monta
    a URL pública (`/api/files/docx/...`). O documento é montado em memória e só é
    escrito no final (`save`), então falhas durante a montagem não deixam arquivo
    parcial.
    """

    def __init__(
        self,
        *,
        output_dir: Path | None = None,
        url_prefix: str = "/api/files",
    ) -> None:
        """Configura o destino e o prefixo de URL do writer."""
        self._output = DocumentOutput("docx", output_dir=output_dir, url_prefix=url_prefix)

    async def write(self, spec: DocumentSpec) -> DocumentResult:
        """Gera o `.docx` do `spec` (fora do event loop) e retorna o resultado."""
        if not isinstance(spec, DocxSpec):
            raise TypeError("DocxWriter só aceita DocxSpec.")
        return await asyncio.to_thread(self._write_sync, spec)

    def _write_sync(self, spec: DocxSpec) -> DocumentResult:
        document = Document()
        document.add_heading(spec.title, level=0)
        for block in spec.blocks:
            self._render_block(document, block)

        path, url = self._output.allocate(spec.extension)
        document.save(str(path))
        return DocumentResult(path=str(path), url=url, metadata=spec.metadata())

    def _render_block(self, document: Document, block: object) -> None:
        if isinstance(block, Heading):
            document.add_heading(block.text, level=block.level)
        elif isinstance(block, Paragraph):
            document.add_paragraph(block.text)
        elif isinstance(block, ListBlock):
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                document.add_paragraph(item, style=style)
        elif isinstance(block, Table):
            self._render_table(document, block)
        elif isinstance(block, ImageRef):
            self._render_image(document, block)

    @staticmethod
    def _render_table(document: Document, block: Table) -> None:
        table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
        table.style = "Table Grid"
        for row_idx, row in enumerate(block.rows):
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = value
                if block.header and row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    @staticmethod
    def _render_image(document: Document, block: ImageRef) -> None:
        if not Path(block.path).is_file():
            raise RuntimeError(f"Imagem não encontrada: {block.path!r}.")
        if block.width_inches is not None:
            document.add_picture(block.path, width=Inches(block.width_inches))
        else:
            document.add_picture(block.path)
