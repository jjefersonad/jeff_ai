"""Converter semântico HTML → DOCX via python-docx (stdlib HTMLParser).

Escolha do spike (html-document-tools-task-docx-spike-1):
- Preferimos um **módulo próprio** em vez de `htmldocx` (descontinuado) ou
  `html-for-docx` (fork ok, mas traz BeautifulSoup e estilos print-CSS
  incompletos). O design pede conversão **semântica** (heading/p/table),
  não fidelidade CSS de impressão — python-docx já está no projeto.
- CSS é ignorado nesta v1 (mesmo contrato do port: aceito, não aplicado).
"""
from __future__ import annotations

from html.parser import HTMLParser
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.table import Table

from src.application.ports.html_document_converter import HtmlDocumentConverter
from src.domain.shared.errors import DomainError

_HEADING_TAGS = {f"h{i}": i for i in range(1, 7)}
# Tags com conteúdo a ignorar (profundidade +1/-1 no start/end).
_SKIP_TAGS = frozenset({"script", "style", "head", "title"})
# Void tags: ignorar sem alterar skip_depth (HTMLParser não emite endtag).
_VOID_SKIP_TAGS = frozenset({"meta", "link"})
_BLOCK_TAGS = frozenset({"p", "div", "li", "blockquote", "pre", "section", "article"})


class _HtmlToDocxBuilder(HTMLParser):
    """Walks HTML and appends python-docx blocks (heading, paragraph, table)."""

    def __init__(self, document: Document) -> None:
        super().__init__(convert_charrefs=True)
        self._doc = document
        self._skip_depth = 0
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] | None = None
        self._cell_parts: list[str] = []
        self._in_cell = False
        self._block_parts: list[str] = []
        self._block_tag: str | None = None
        self._list_ordered = False
        self._list_index = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        lower = tag.lower()
        if lower in _VOID_SKIP_TAGS:
            return
        if lower in _SKIP_TAGS:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return

        if lower == "table":
            self._flush_block()
            self._in_table = True
            self._table_rows = []
            return
        if lower == "tr" and self._in_table:
            self._current_row = []
            return
        if lower in {"td", "th"} and self._in_table:
            self._in_cell = True
            self._cell_parts = []
            return
        if lower in _HEADING_TAGS:
            self._flush_block()
            self._block_tag = lower
            self._block_parts = []
            return
        if lower in {"ul", "ol"}:
            self._flush_block()
            self._list_ordered = lower == "ol"
            self._list_index = 0
            return
        if lower == "li":
            self._flush_block()
            self._block_tag = "li"
            self._block_parts = []
            self._list_index += 1
            return
        if lower in _BLOCK_TAGS:
            self._flush_block()
            self._block_tag = "p"
            self._block_parts = []
            return
        if lower == "br":
            if self._in_cell:
                self._cell_parts.append("\n")
            elif self._block_tag:
                self._block_parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        lower = tag.lower()
        if lower in _SKIP_TAGS:
            if self._skip_depth:
                self._skip_depth -= 1
            return
        if self._skip_depth:
            return

        if lower in {"td", "th"} and self._in_table and self._in_cell:
            text = "".join(self._cell_parts).strip()
            if self._current_row is not None:
                self._current_row.append(text)
            self._in_cell = False
            self._cell_parts = []
            return
        if lower == "tr" and self._in_table and self._current_row is not None:
            self._table_rows.append(self._current_row)
            self._current_row = None
            return
        if lower == "table" and self._in_table:
            self._emit_table(self._table_rows)
            self._in_table = False
            self._table_rows = []
            return
        if lower in _HEADING_TAGS or lower in _BLOCK_TAGS or lower == "li":
            self._flush_block()
            return
        if lower in {"ul", "ol"}:
            self._list_ordered = False
            self._list_index = 0

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if self._in_cell:
            self._cell_parts.append(data)
            return
        if self._block_tag is not None:
            self._block_parts.append(data)
            return
        # Texto solto fora de bloco → parágrafo.
        text = data.strip()
        if text:
            self._block_tag = "p"
            self._block_parts = [text]
            self._flush_block()

    def close(self) -> None:
        self._flush_block()
        super().close()

    def _flush_block(self) -> None:
        if self._block_tag is None:
            return
        text = "".join(self._block_parts).strip()
        tag = self._block_tag
        self._block_tag = None
        self._block_parts = []
        if not text:
            return
        if tag in _HEADING_TAGS:
            level = _HEADING_TAGS[tag]
            self._doc.add_heading(text, level=level)
            return
        if tag == "li":
            prefix = f"{self._list_index}. " if self._list_ordered else "• "
            para = self._doc.add_paragraph(prefix + text)
            para.paragraph_format.left_indent = Pt(18)
            return
        self._doc.add_paragraph(text)

    def _emit_table(self, rows: list[list[str]]) -> None:
        if not rows:
            return
        cols = max(len(r) for r in rows)
        if cols == 0:
            return
        table: Table = self._doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx in range(cols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                table.cell(r_idx, c_idx).text = cell_text


class HtmlDocxConverter(HtmlDocumentConverter):
    """HTML semântico → bytes DOCX (python-docx)."""

    async def convert(self, *, html: str, css: str | None, kind: str) -> bytes:
        # `css` intencionalmente ignorado — conversão semântica, não print-CSS.
        _ = css
        if kind != "docx":
            raise DomainError(
                f"HtmlDocxConverter só suporta kind='docx', recebeu {kind!r}."
            )
        if not html or not html.strip():
            raise DomainError("HTML vazio: nada para converter para DOCX.")

        document = Document()
        builder = _HtmlToDocxBuilder(document)
        builder.feed(html)
        builder.close()

        # Evita documento completamente vazio (só estilos default).
        has_content = bool(document.paragraphs) or bool(document.tables)
        if not has_content:
            raise DomainError("HTML sem conteúdo convertível para DOCX.")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
